import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.dml.color import RGBColor

# Path
md_path = r"D:\Gacoan\Tugas_Data_Science\Makalah_Data_Science.md"
docx_path = r"D:\Gacoan\Tugas_Data_Science\Makalah_Data_Science.docx"

print("Mengubah Markdown ke DOCX...")

doc = Document()

# Page setup (margins)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Default style formatting
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Read Markdown
with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

in_table = False
table_headers = []
table_rows = []

# Regex helper to remove markdown bold/italic inside table cells
def clean_md_formatting(text):
    text = re.sub(r'\*\*(.*?)\*\*|__(.*?)__', r'\1\2', text)
    text = re.sub(r'\*(.*?)\*|_(.*?)_', r'\1\2', text)
    return text.strip()

for line in lines:
    line_strip = line.strip()
    
    # Handle Tables
    if line_strip.startswith('|'):
        in_table = True
        cells = [clean_md_formatting(c) for c in line_strip.split('|')[1:-1]]
        if '---' in line_strip:
            # Separator line, ignore
            continue
        if not table_headers:
            table_headers = cells
        else:
            table_rows.append(cells)
        continue
    else:
        # If we were in a table and it ended, construct and write the table in docx
        if in_table and table_headers:
            # Create table
            table = doc.add_table(rows=1, cols=len(table_headers))
            table.style = 'Light Shading Accent 1'
            hdr_cells = table.rows[0].cells
            for idx, name in enumerate(table_headers):
                hdr_cells[idx].text = name
                hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
            
            for row_data in table_rows:
                row_cells = table.add_row().cells
                for idx, cell_val in enumerate(row_data):
                    if idx < len(row_cells):
                        row_cells[idx].text = cell_val
            
            doc.add_paragraph() # Add empty spacing after table
            table_headers = []
            table_rows = []
            in_table = False
            
    # Handle Headings
    if line_strip.startswith('# '):
        p = doc.add_paragraph()
        run = p.add_run(line_strip[2:])
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(41, 128, 185)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
    elif line_strip.startswith('## '):
        p = doc.add_paragraph()
        run = p.add_run(line_strip[3:])
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(44, 62, 80)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif line_strip.startswith('### '):
        p = doc.add_paragraph()
        run = p.add_run(line_strip[4:])
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(44, 62, 80)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif line_strip.startswith('#### '):
        p = doc.add_paragraph()
        run = p.add_run(line_strip[5:])
        run.font.size = Pt(11)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        
    # Handle Separator Line
    elif line_strip == '---':
        p = doc.add_paragraph()
        p.add_run("―" * 60).font.color.rgb = RGBColor(189, 195, 199)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    # Handle Bullet Points
    elif line_strip.startswith('- ') or line_strip.startswith('* '):
        p = doc.add_paragraph(style='List Bullet')
        # Check bold formatting inside bullet
        content = line_strip[2:]
        bold_match = re.search(r'\*\*(.*?)\*\*', content)
        if bold_match:
            bold_text = bold_match.group(1)
            normal_text = content.replace(f"**{bold_text}**", "")
            
            run_bold = p.add_run(bold_text)
            run_bold.bold = True
            p.add_run(normal_text)
        else:
            p.add_run(content)
            
    # Handle Numbered Lists
    elif re.match(r'^\d+\.', line_strip):
        p = doc.add_paragraph(style='List Number')
        content = re.sub(r'^\d+\.\s*', '', line_strip)
        
        bold_match = re.search(r'\*\*(.*?)\*\*', content)
        if bold_match:
            bold_text = bold_match.group(1)
            normal_text = content.replace(f"**{bold_text}**", "")
            
            run_bold = p.add_run(bold_text)
            run_bold.bold = True
            p.add_run(normal_text)
        else:
            p.add_run(content)
            
    # Handle Blockquotes / ASCI Art
    elif line_strip.startswith('│') or line_strip.startswith('┌') or line_strip.startswith('└') or line_strip.startswith('├'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        p.paragraph_format.left_indent = Inches(0.5)
        
    # Handle Normal Paragraphs
    elif line_strip:
        # Check if line is formula
        if line_strip.startswith('$$') and line_strip.endswith('$$'):
            p = doc.add_paragraph()
            run = p.add_run(line_strip[2:-2])
            run.font.name = 'Cambria Math'
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p = doc.add_paragraph()
            # Process inline formatting like **bold**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            p.paragraph_format.space_after = Pt(6)

doc.save(docx_path)
print(f"File Word berhasil disimpan di: {docx_path}")
